import os
import json
import re
import docx
import fitz # PyMuPDF
from odf.opendocument import load as load_odf
from odf import text as odf_text

DOCS_DIR = "/Users/gustavoalmeida/Cartorio/documents"
OUTPUT_JSON = "/Users/gustavoalmeida/Cartorio/inventory.json"

def read_txt(filepath):
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()

def read_docx(filepath):
    doc = docx.Document(filepath)
    fullText = []
    for para in doc.paragraphs:
        if para.text:
            fullText.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                fullText.append(" | ".join(row_text))
    return "\n".join(fullText)

def read_pdf(filepath):
    doc = fitz.open(filepath)
    text = []
    for page in doc:
        text.append(page.get_text())
    return "\n".join(text)

def read_odt(filepath):
    doc = load_odf(filepath)
    paragraphs = doc.getElementsByType(odf_text.P)
    res = []
    for p in paragraphs:
        t_list = []
        for child in p.childNodes:
            if hasattr(child, 'data'):
                t_list.append(str(child.data))
        if t_list:
            res.append("".join(t_list))
    return "\n".join(res)

def extract_file_data(filename):
    filepath = os.path.join(DOCS_DIR, filename)
    ext = os.path.splitext(filename)[1].lower()
    text = ""
    try:
        if ext == ".docx":
            text = read_docx(filepath)
        elif ext == ".pdf":
            text = read_pdf(filepath)
        elif ext == ".txt":
            text = read_txt(filepath)
        elif ext == ".odt":
            text = read_odt(filepath)
    except Exception as e:
        print(f"Error reading {filename}: {e}")
        text = f"[ERROR READING FILE: {e}]"
    
    word_count = len(text.split())
    char_count = len(text)
    
    category = "Geral / Outros"
    fname_upper = filename.upper()
    text_upper = text.upper()[:1000]
    
    if "TESTAMENTO" in fname_upper or "TESTAMEN" in text_upper:
        category = "Testamento"
    elif "USUCAPI" in fname_upper or "USUCAPI" in text_upper:
        category = "Usucapião"
    elif "INVENTÁRIO" in fname_upper or "INVENTARIO" in fname_upper or "INVENTÁRIO" in text_upper or "INVENTARIO" in text_upper:
        category = "Inventário e Partilha"
    elif "DIVÓRCIO" in fname_upper or "DIVORCIO" in fname_upper or "DIVÓRCIO" in text_upper or "DIVORCIO" in text_upper:
        category = "Divórcio e Separação"
    elif "UNIÃO ESTÁVEL" in fname_upper or "UNIAO ESTAVEL" in fname_upper or "UNIÃO ESTÁVEL" in text_upper or "UNIAO ESTAVEL" in text_upper:
        category = "União Estável e Casamento"
    elif "TABELA" in fname_upper or "EMOLUMENTO" in text_upper:
        category = "Tabelas e Emolumentos"
    elif "PROVIMENTO" in fname_upper or "CNJ" in text_upper or "JURISPRUDÊNCIA" in text_upper or "JURIS" in fname_upper:
        category = "Normas e Provimentos CNJ"
    elif "ADJUDICAÇÃO" in fname_upper or "ADJUDICACAO" in fname_upper or "ADJUDICAÇÃO" in text_upper:
        category = "Adjudicação Compulsória"
    elif "ESTREMAÇÃO" in fname_upper or "ESTREMACAO" in fname_upper or "ESTREMAÇÃO" in text_upper:
        category = "Estremação"
    elif "APOSTILAMENTO" in fname_upper or "ATA NOTARIAL" in fname_upper or "FIRM" in fname_upper or "PROCURAÇÃO" in fname_upper or "PROCURACAO" in fname_upper:
        category = "Atos Notariais Diversos"
    
    return {
        "filename": filename,
        "extension": ext,
        "size_bytes": os.path.getsize(filepath),
        "word_count": word_count,
        "char_count": char_count,
        "initial_category": category,
        "text": text
    }

def main():
    files = sorted(os.listdir(DOCS_DIR))
    inventory = []
    print(f"Processing {len(files)} files...")
    for f in files:
        if f.startswith('.'): continue
        data = extract_file_data(f)
        inventory.append(data)
    
    with open(OUTPUT_JSON, "w", encoding="utf-8") as out:
        json.dump(inventory, out, ensure_ascii=False, indent=2)
    
    print(f"\nInventory successfully generated at {OUTPUT_JSON} with {len(inventory)} documents.")

if __name__ == "__main__":
    main()
